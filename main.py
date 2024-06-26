import streamlit as st
from docx import Document
from io import BytesIO
import os
import base64
import mammoth
import openai
import json
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import copy
from copy import deepcopy

# Connect to OpenAI key
openai.api_key = st.secrets["OPENAI_API_KEY"]

# Function to load a DOCX file
def load_docx(document):
    doc = Document(document)
    return doc

def copy_document(doc):
    new_doc = Document()
    for element in doc.element.body:
        new_doc.element.body.append(copy.deepcopy(element))
    return new_doc

# Function to generate content using OpenAI and replace targeted text
def generate_content(doc, nl_instruction):
    instruction_message = f"You are a professional business analyst. Your task is to generate a report based on the user's requirements. Do not use markdown format for your response."

    openai_response = openai.chat.completions.create(
        model="gpt-4-turbo-preview",
        messages=[
            {"role": "system", "content": instruction_message},
            {"role": "user", "content": nl_instruction}
        ],
        temperature=0.5,
        stream=False,
    )
    generated_text = openai_response.choices[0].message.content

    # Create a copy of the document to avoid modifying the original
    modified_doc = deepcopy(doc)

    for paragraph in modified_doc.paragraphs:
        if '[targeted_text]' in paragraph.text:
            paragraph.text = generated_text

    return modified_doc

def generate_table(doc, nl_instruction):
    instruction_message = "Generate a table with sections: User Story, Pre-condition, Description with User Workflow, Post-condition, Acceptance Criteria, and Edge Case."
    instruction_message += " Your output should return a list of tuples where each tuple contains the section name and description."
    instruction_message += """ Here's an example json: {
        "descriptions": [
            {
            "section": "User Story",
            "description": "As a user, I want to securely log in to the web application so that I can access my personal dashboard."
            },
            {
            "section": "Pre-condition",
            "description": "The user must be registered with the web application, having a valid username and password."
            }]
            }
        """

    openai_response = openai.chat.completions.create(
        model="gpt-4-turbo-preview",
        response_format={"type": "json_object"},
        messages=[
            {"role": "system", "content": instruction_message},
            {"role": "user", "content": nl_instruction}
        ],
        temperature=0.5,
        stream=False,
    )
    descriptions = openai_response.choices[0].message.content
    try:
        results = json.loads(descriptions)
        descriptions = results['descriptions']
    except json.JSONDecodeError:
        results = {}
        descriptions = []

    # Create a new document
    new_doc = Document()

    # Iterate over the original document
    for para in doc.paragraphs:
        # If the paragraph is the placeholder, add the table
        if '[targeted_table]' in para.text:
            # Create a table
            table = new_doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "Section Name"
            hdr_cells[1].text = "Section Description"
            # Set column widths
            hdr_cells[0].width = Inches(2)
            hdr_cells[1].width = Inches(4)
            # Center align the header text and make it bold
            for idx, header in enumerate(hdr_cells):
                header.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in header.paragraphs[0].runs:
                    run.bold = True

            # Add rows to the table
            for item in descriptions:
                row_cells = table.add_row().cells
                row_cells[0].text = item['section']
                row_cells[1].text = item['description']
        else:
            # If the paragraph is not the placeholder, just copy it
            new_doc.add_paragraph(para.text)

    return new_doc

# Function to save a DOCX file to a buffer
def save_docx(doc):
    buffer = BytesIO()
    doc.save(buffer)
    return buffer

# Function to convert DOCX to HTML
def docx_to_html(docx_path, html_path):
    try:
        with open(docx_path, "rb") as docx_file:
            result = mammoth.convert_to_html(docx_file)
            html = result.value

        with open(html_path, "w", encoding="utf-8") as html_file:
            html_file.write(html)
        return True
    except Exception as e:
        print(f"Failed to convert DOCX to HTML: {e}")
        return False

def main():
    st.title('Agent Jude - Requirements Analyst')

    # Initialize session state for documents
    if 'product_doc' not in st.session_state:
        st.session_state.product_doc = None
    if 'modified_doc' not in st.session_state:
        st.session_state.modified_doc = None

    temp_dir = "./temp"
    os.makedirs(temp_dir, exist_ok=True)

    uploaded_file = st.file_uploader("Step 1. Choose a .docx file", type="docx")
    if uploaded_file is not None:
        document = load_docx(uploaded_file)

        nl_instruction = st.text_area("Step 2. Describe your product requirements:", placeholder="Create a simple requirements for a login page of a web application...")

        modified_doc = None  # Initialize modified_doc here to prevent UnboundLocalError

        if st.button("Generate product documentation", help="Generate doc at [target_text]", type="primary"):
            st.session_state.product_doc = generate_content(document, nl_instruction)

            # Display the preview of the modified document (product documentation)
            buffer = save_docx(st.session_state.product_doc)
            buffer.seek(0)

            temp_docx_path = os.path.join(temp_dir, 'modified_document.docx')
            with open(temp_docx_path, 'wb') as f:
                f.write(buffer.getvalue())

            html_path = os.path.join(temp_dir, 'temporary_copy.html')

            if docx_to_html(temp_docx_path, html_path):
                with open(html_path, "r", encoding="utf-8") as html_file:
                    html_content = html_file.read()
                    st.subheader("✅ Modified Document Preview (Product Documentation):", divider="green")
                    st.markdown("""---""")
                    st.markdown(html_content, unsafe_allow_html=True)
                    st.download_button(
                        label="Download modified DOCX (Product Documentation)",
                        data=buffer,
                        file_name="modified_document_product_doc.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        type="primary"
                    )

        if st.button("Generate user story table", help="Generate a table at the [target_table]", type="primary"):
            temp_docx_path = os.path.join(temp_dir, 'modified_document.docx')
            if temp_docx_path:  # Check if product_doc has been created
                table_doc = deepcopy(st.session_state.product_doc)
                st.session_state.modified_doc = generate_table(table_doc, nl_instruction)

                # Save the modified document to a buffer for download
                buffer = save_docx(st.session_state.modified_doc)
                buffer.seek(0)
                
                # Save the modified DOCX to a file in the temporary directory
                temp_docx_path = os.path.join(temp_dir, 'modified_document_with_table.docx')
                with open(temp_docx_path, 'wb') as f:
                    f.write(buffer.getvalue())
                
                # Create an HTML path in the temporary directory for preview
                html_path = os.path.join(temp_dir, 'temporary_copy_with_table.html')
                
                # Convert the DOCX to HTML for preview
                if docx_to_html(temp_docx_path, html_path):
                    with open(html_path, "r", encoding="utf-8") as html_file:
                        html_content = html_file.read()
                        st.subheader("✅ Modified Document Preview (User Story Table):", divider="green")
                        st.markdown("""---""")
                        st.markdown(html_content, unsafe_allow_html=True)
                        st.download_button(
                            label="Download modified DOCX (User Story Table)",
                            data=buffer,
                            file_name="modified_document_with_table.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            type="primary"
                        )
                else:
                    st.error("Failed to convert the document to HTML.")
            else:
                st.error("Please generate the product documentation first.")

        # Save the uploaded DOCX to a temporary file
        temp_uploaded_docx_path = os.path.join(temp_dir, 'uploaded_document.docx')
        with open(temp_uploaded_docx_path, 'wb') as f:
            f.write(uploaded_file.getvalue())

        # Convert the uploaded DOCX to HTML for preview
        temp_html_path = os.path.join(temp_dir, 'uploaded_copy.html')
        if docx_to_html(temp_uploaded_docx_path, temp_html_path):
            with open(temp_html_path, "r", encoding="utf-8") as html_file:
                html_content = html_file.read()
                st.subheader("Uploaded Document Preview:",divider='rainbow')
                st.markdown("""---""")
                st.markdown(html_content, unsafe_allow_html=True)
        else:
            st.error("Failed to convert the uploaded document to HTML.")
            
            # Save the modified DOCX to a file in the temp directory
            temp_docx_path = os.path.join(temp_dir, 'modified_document.docx')
            with open(temp_docx_path, 'wb') as f:
                f.write(buffer.getvalue())
            
            # Create a HTML path in the temp directory
            html_path = os.path.join(temp_dir, 'temporary_copy.html')
            
            # Convert DOCX to HTML
            if docx_to_html(temp_docx_path, html_path) and st.session_state.product_doc is not None:
                with open(html_path, "r", encoding="utf-8") as html_file:
                    html_content = html_file.read()
                    st.subheader("✅ Modified Document Preview (User Story Table):", divider="green")
                    st.markdown("""---""")
                    st.markdown(html_content, unsafe_allow_html=True)
                    st.download_button(
                        label="Download modified DOCX (User Story Table)",
                        data=buffer,
                        file_name="modified_document_with_table.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        type="primary"
                    )
            else:
                st.error("Failed to convert the document to HTML.")

if __name__ == "__main__":
    main()