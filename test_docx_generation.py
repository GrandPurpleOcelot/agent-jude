from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime

document = Document()

document.add_picture('FPT_Software_Logo.png')
document.add_heading('PRODUCT REQUIREMENTS', 0)
# Add a 2x2 table with version and date
table = document.add_table(rows=2, cols=2)
table.style = 'Table Grid'

# Add headers
headers = ['Version', 'Date']
for i, header in enumerate(headers):
    table.cell(0, i).text = header

# Add data
version = '1.0'
date = datetime.date.today().strftime('%Y-%m-%d')
table.cell(1, 0).text = version
table.cell(1, 1).text = date
document.add_page_break()

document.add_heading('Example Document Title', 0)

p = document.add_paragraph('A plain paragraph having some text')
p.add_run('bold').bold = True
p.add_run(' and some ')
p.add_run('italic.').italic = True

document.add_heading('Heading, level 1', level=1)
document.add_paragraph('Intense quote', style='Intense Quote')

document.add_paragraph(
    'first item in unordered list', style='List Bullet'
)
document.add_paragraph(
    'first item in ordered list', style='List Number'
)

document.add_paragraph(
    '[targeted_text]', style='Normal'
)

document.add_heading('US01_Login Function', level=2)
document.add_paragraph(
    '[targeted_table]', style='Normal'
)

# document.add_picture('example_image.png', width=Inches(1.25))

# document.add_page_break()

# document.add_heading('US02_Filter Function', level=2)
# # Create a table with the necessary rows and columns
# # Here we need 7 rows and 2 columns
# table = document.add_table(rows=7, cols=2)
# table.style = 'Table Grid'

# # Define custom widths for the columns
# widths = [Inches(1), Inches(5)]  # First column is 1 inches, second is 5 inches

# # Apply widths to each cell in the columns
# for row in table.rows:
#     row.cells[0].width = widths[0]
#     row.cells[1].width = widths[1]

# # Define the headers and contents
# headers = ["Section Name", "Section Description"]
# descriptions = [
#     ("User Story", "As a user, I want to be able to filter the data in the table based on specific criteria, so that I can quickly find and view the relevant information."),
#     ("Pre-condition", "- The user is on the page displaying the table data.\n- The table data contains multiple columns."),
#     ("Description with User Workflow", "1. The user will see a filter input field or dropdown menu above or beside the table.\n2. The user can enter a search term or select a filter condition from the dropdown menu.\n3. The filter condition can be based on one or more columns in the table (e.g., filter by name, date, status, etc.).\n4. As the user types or selects a filter condition, the table data will automatically update and display only the rows that match the filter criteria.\n5. If no rows match the filter criteria, the table will display a message indicating that no data is available.\n6. The user can clear the filter by removing the search term or resetting the filter condition."),
#     ("Post-condition", "- The table displays only the data that matches the specified filter criteria.\n- If no data matches the filter criteria, the table is empty with a message indicating that no data is available."),
#     ("Acceptance Criteria", "1. The filter feature should work for all columns in the table.\n2. The filter should support partial matches (e.g., typing 'John' should display rows containing 'John Doe', 'Johnson', etc.).\n3. The filter should be case-insensitive (e.g., typing 'JOHN' should match 'John Doe').\n4. The filter should update the table data in real-time as the user types or changes the filter condition.\n5. The filter should be clearable, allowing the user to reset the table to display all data.\n6. The filter should handle empty or null values in the table data.\n7. The filter should provide a visual indication of the applied filter criteria (e.g., highlighting the filter input field, displaying the filter condition, etc.)."),
#     ("Edge Case", "1. If the table data contains a large number of rows (e.g., thousands or millions of rows), the filter should be optimized to handle the large dataset efficiently and provide a smooth user experience.\n2. If the table data contains complex data types (e.g., dates, currencies, or custom formats), the filter should be able to handle those data types correctly.\n3. If the table data contains nested or hierarchical data structures, the filter should be able to handle those structures and provide a way to filter at different levels of the hierarchy.")
# ]

# # Fill in the table header
# for i, header in enumerate(headers):
#     table.cell(0, i).text = header

# # Fill in the table data
# for i, (title, desc) in enumerate(descriptions, start=1):
#     table.cell(i, 0).text = title
#     table.cell(i, 1).text = desc

# # Set column headers
# hdr_cells = table.rows[0].cells
# for idx, header in enumerate(headers):
#     hdr_cells[idx].text = header
#     hdr_cells[idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER  # Center align the header text
#     for run in hdr_cells[idx].paragraphs[0].runs:
#         run.bold = True  # Make the header text bold

document.save('demo.docx')